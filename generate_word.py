import os
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 讀取配置文件
with open('config.json', 'r') as config_file:
    config = json.load(config_file)

image_base_folder = config['image_base_folder']
output_file = config['output_file']
images_per_row = config['images_per_row']
images_per_column = config['images_per_column']
image_spacing = config['image_spacing']
page_height = config['page_height']
page_width = config['page_width']
page_border_width = config['page_border_width']

# 計算每頁最多圖片數量
max_images_per_page = images_per_row * images_per_column

# 建立 Word 文件
doc = Document()

# 設置頁邊距
sections = doc.sections
for section in sections:
    section.top_margin = Inches(page_border_width)
    section.bottom_margin = Inches(page_border_width)
    section.left_margin = Inches(page_border_width)
    section.right_margin = Inches(page_border_width)

# 計算本週日期範圍，從上週日開始
today = datetime.today()
start_of_week = today - timedelta(days=(today.weekday() + 1 if today.weekday() != 6 else 0))
end_of_week = start_of_week + timedelta(days=6)
date_range_str = start_of_week.strftime('%m/%d') + ' ~ ' + end_of_week.strftime('%m/%d')

# 函數：隱藏內部框線，只保留外部框線
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # 隱藏內部邊框
    for border_name in ['insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'nil')

    # 設置外層邊框
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:space'), '0')
        border.set(qn('w:sz'), '4')  # 邊框寬度

# 函數：插入圖片框架
def add_image_frame(doc, title, image_paths, bAddpagebreak):
    num_images = len(image_paths)
    num_pages = (num_images + max_images_per_page - 1) // max_images_per_page  # 計算總頁數

    if (bAddpagebreak):
            doc.add_page_break()
    

    for page in range(num_pages):
        # 插入標題和本週日期範圍到同一行
        header_table = doc.add_table(rows=1, cols=2)
        #header_table.autofit = True
        header_cells = header_table.rows[0].cells

        # 設置左側標題單元格
        title_cell = header_cells[0]
        run = title_cell.paragraphs[0].add_run(title)
        # 設置標題單元格的段落格式
        title_cell.paragraphs[0].paragraph_format.space_after = Inches(0.01)  # 調整段落後間距
        title_cell.paragraphs[0].paragraph_format.space_before = Inches(0.01)  # 調整段落前間距        
        run.bold = True
        run.font.size = Inches(0.3)


        # 設置右側日期範圍單元格
        date_cell = header_cells[1]
        # 設置日期單元格的段落格式
        date_cell.paragraphs[0].paragraph_format.space_after = Inches(0.01)  # 調整段落後間距
        date_cell.paragraphs[0].paragraph_format.space_before = Inches(0.01)  # 調整段落前間距
                
        run = date_cell.paragraphs[0].add_run(date_range_str)
        run.bold = True
        run.font.size = Inches(0.3)
        date_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 插入一个空行并设置行高最小
        empty_paragraph = doc.add_paragraph()
        empty_paragraph.paragraph_format.space_after = Inches(0)  # 段落后间距
        empty_paragraph.paragraph_format.space_before = Inches(0)  # 段落前间距
        empty_paragraph.paragraph_format.line_spacing = Inches(0.1)  # 行间距


        start_index = page * max_images_per_page
        end_index = min(start_index + max_images_per_page, num_images)  # 確保 end_index 不超出範圍
        batch_images = image_paths[start_index:end_index]

        # 創建表格來排版圖片，占據整頁高度
        table = doc.add_table(rows=images_per_column, cols=images_per_row)
        table.autofit = True
        set_table_borders(table)

        # 計算單元格的寬度和高度
        cell_width = (page_width - image_spacing * (images_per_row - 1)) / images_per_row
        cell_width_emu = Inches(cell_width)
        cell_height = Inches((page_height - page_border_width - image_spacing * (images_per_column - 1)) / images_per_column)        

        # 設置單元格的尺寸和對齊方式
        for row in table.rows:
            row.height = cell_height  # 設置行高
            for cell in row.cells:
                cell.width = cell_width_emu  # 調整單元格寬度
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 插入圖片到表格
        for idx, image_path in enumerate(batch_images):
            row_idx = idx // images_per_row
            col_idx = idx % images_per_row

            # 確保索引在範圍內
            if row_idx < images_per_column and col_idx < images_per_row:
                cell = table.cell(row_idx, col_idx)
                
                if os.path.exists(image_path):
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(image_path, width=cell_width_emu)
                else:
                    cell.text = "圖片未找到： " + image_path

        # 如果還有剩餘圖片，添加分頁
        if page < num_pages - 1:
            doc.add_page_break()

# 獲取子資料夾列表
subfolders = [f.path for f in os.scandir(image_base_folder) if f.is_dir()]

bAddpagebreak = False

# 遍歷每個子資料夾
for folder in subfolders:
    # 獲取圖片列表
    image_files = [os.path.join(folder, f) for f in os.listdir(folder) if f.endswith(('.jpg', '.png', '.jpeg'))]
    
    # 插入圖片框架
    add_image_frame(doc, os.path.basename(folder), image_files, bAddpagebreak)
    bAddpagebreak = True


# 保存 Word 文件
doc.save(output_file)
